import os
import sys
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("需要安装 python-docx 库：")
    print("  pip install python-docx")
    sys.exit(1)

EXCLUDE_DIRS = {
    'node_modules', '.git', '.svn', '__pycache__', '.pytest_cache',
    '.mypy_cache', '.tox', '.venv', 'venv', 'env', '.env',
    'dist', 'build', 'out', '.next', '.nuxt', '.cache', 'coverage',
    '.idea', '.vscode', 'target', 'bin', 'obj', '.gradle',
    '.dart_tool', '.pub-cache', 'Pods', 'DerivedData', '.ai-agent-temp',
}

EXCLUDE_EXTENSIONS = {
    '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.svg', '.webp',
    '.tif', '.tiff', '.psd', '.ai', '.eps', '.raw', '.cr2', '.nef',
    '.heic', '.heif', '.avif',
    '.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm', '.m4v',
    '.mpg', '.mpeg', '.3gp', '.ogv',
    '.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a', '.opus',
    '.mid', '.midi',
    '.ttf', '.otf', '.woff', '.woff2', '.eot',
    '.zip', '.tar', '.gz', '.bz2', '.7z', '.rar', '.xz', '.zst', '.tgz',
    '.exe', '.dll', '.so', '.dylib', '.o', '.a', '.lib',
    '.pyc', '.pyo', '.class', '.jar', '.war', '.wasm', '.map',
    '.db', '.sqlite', '.sqlite3', '.mdb',
    '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
    '.bin', '.dat', '.iso', '.img', '.dmg', '.vsix', '.lock',
}

EXCLUDE_FILES = {
    '.DS_Store', 'Thumbs.db', 'desktop.ini',
    'package-lock.json', 'yarn.lock', 'pnpm-lock.yaml',
    'Cargo.lock', 'Gemfile.lock', 'poetry.lock', 'composer.lock', 'go.sum',
}

MAX_FILE_SIZE = 200 * 1024

# 用户通过 --exclude 指定的排除列表
USER_EXCLUDES = []


def is_text_file(filepath):
    try:
        with open(filepath, 'rb') as f:
            chunk = f.read(8192)
            if b'\x00' in chunk:
                return False
        return True
    except (IOError, PermissionError):
        return False


def matches_user_exclude(relpath, filename):
    """检查文件是否匹配用户指定的排除规则

    支持的排除格式：
      - 文件名:        export_code.py
      - 相对路径:      edge-extension/popup.js
      - 目录名:        edge-extension/
      - 通配符后缀:    *.md
      - 通配符前缀:    test_*
    """
    normalized = relpath.replace('\\', '/')

    for pattern in USER_EXCLUDES:
        p = pattern.replace('\\', '/').strip()
        if not p:
            continue

        # 排除整个目录: "edge-extension/" 或 "edge-extension"
        dir_pattern = p.rstrip('/')
        if normalized.startswith(dir_pattern + '/') or normalized == dir_pattern:
            return True

        # 精确匹配文件名: "popup.js"
        if filename == p:
            return True

        # 精确匹配相对路径: "edge-extension/popup.js"
        if normalized == p:
            return True

        # 通配符后缀: "*.md"
        if p.startswith('*.'):
            ext_pattern = p[1:]  # ".md"
            if filename.endswith(ext_pattern):
                return True

        # 通配符前缀: "test_*"
        if p.endswith('*'):
            prefix = p[:-1]
            if filename.startswith(prefix):
                return True

        # 路径中包含通配符目录: "test/*"
        if '/*' == p[-2:] and normalized.startswith(p[:-2] + '/'):
            return True

    return False


def should_include(filepath, filename, relpath):
    if filename in EXCLUDE_FILES:
        return False
    _, ext = os.path.splitext(filename)
    if ext.lower() in EXCLUDE_EXTENSIONS:
        return False
    if matches_user_exclude(relpath, filename):
        return False
    try:
        size = os.path.getsize(filepath)
        if size > MAX_FILE_SIZE or size == 0:
            return False
    except OSError:
        return False
    if not is_text_file(filepath):
        return False
    return True


def strip_comments(content, ext):
    if ext not in ('.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.go',
                   '.c', '.cpp', '.h', '.hpp', '.cs', '.rs', '.rb',
                   '.php', '.swift', '.kt', '.scala', '.sh', '.bash',
                   '.ps1', '.r', '.pl', '.lua'):
        return content

    lines = content.split('\n')
    result = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            result.append(line)
            continue

        has_chinese = any('\u4e00' <= ch <= '\u9fff' for ch in stripped)

        if not has_chinese:
            result.append(line)
            continue

        if ext in ('.py', '.rb', '.sh', '.bash', '.ps1', '.r', '.pl'):
            if stripped.startswith('#'):
                continue
            hash_pos = line.find('#')
            if hash_pos > 0:
                before = line[:hash_pos].rstrip()
                if before:
                    result.append(before)
                    continue
            result.append(line)

        elif ext in ('.js', '.ts', '.jsx', '.tsx', '.java', '.go',
                     '.c', '.cpp', '.h', '.hpp', '.cs', '.rs',
                     '.swift', '.kt', '.scala', '.lua'):
            if stripped.startswith('//'):
                continue
            slash_pos = line.find('//')
            if slash_pos > 0:
                before = line[:slash_pos].rstrip()
                if before:
                    result.append(before)
                    continue
            result.append(line)

        else:
            result.append(line)

    return '\n'.join(result)


def collect_files(root_dir):
    files = []
    for dirpath, dirnames, filenames in os.walk(root_dir):
        dirnames[:] = [
            d for d in dirnames
            if d not in EXCLUDE_DIRS and not d.startswith('.')
        ]
        dirnames.sort()

        for filename in sorted(filenames):
            filepath = os.path.join(dirpath, filename)
            relpath = os.path.relpath(filepath, root_dir)
            if should_include(filepath, filename, relpath):
                files.append((relpath, filepath))
    return files


def get_lang_from_ext(ext):
    mapping = {
        '.py': 'Python', '.js': 'JavaScript', '.ts': 'TypeScript',
        '.jsx': 'JSX', '.tsx': 'TSX', '.java': 'Java',
        '.go': 'Go', '.rs': 'Rust', '.c': 'C', '.cpp': 'C++',
        '.h': 'C/C++ Header', '.hpp': 'C++ Header',
        '.cs': 'C#', '.rb': 'Ruby', '.php': 'PHP',
        '.swift': 'Swift', '.kt': 'Kotlin', '.scala': 'Scala',
        '.sh': 'Shell', '.bash': 'Bash', '.ps1': 'PowerShell',
        '.html': 'HTML', '.css': 'CSS', '.scss': 'SCSS',
        '.json': 'JSON', '.xml': 'XML', '.yaml': 'YAML', '.yml': 'YAML',
        '.md': 'Markdown', '.sql': 'SQL', '.lua': 'Lua',
        '.r': 'R', '.pl': 'Perl', '.bat': 'Batch',
    }
    return mapping.get(ext.lower(), ext.upper().replace('.', ''))


def export_project(root_dir, output_file):
    root_dir = os.path.abspath(root_dir)
    if not os.path.isdir(root_dir):
        print(f"Error: directory not found -> {root_dir}")
        sys.exit(1)

    print(f"Scanning: {root_dir}")

    if USER_EXCLUDES:
        print(f"Excluding: {', '.join(USER_EXCLUDES)}")

    files = collect_files(root_dir)
    print(f"Found {len(files)} code files")

    if not files:
        print("No code files found!")
        sys.exit(1)

    doc = Document()

    # ========== 样式设置 ==========
    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # ========== 标题页 ==========
    title = doc.add_heading('Project Code Export', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_lines = [
        f"Project: {root_dir}",
        f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Files: {len(files)}",
    ]
    if USER_EXCLUDES:
        info_lines.append(f"Excluded: {', '.join(USER_EXCLUDES)}")

    for line in info_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('')

    # ========== 文件目录 ==========
    doc.add_heading('File Index', level=1)

    for i, (relpath, filepath) in enumerate(files, 1):
        size = os.path.getsize(filepath)
        normalized = relpath.replace(os.sep, '/')
        p = doc.add_paragraph()
        run = p.add_run(f"  {i:3d}. {normalized}")
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run_size = p.add_run(f"  ({size:,} bytes)")
        run_size.font.name = 'Consolas'
        run_size.font.size = Pt(8)
        run_size.font.color.rgb = RGBColor(130, 130, 130)

    doc.add_page_break()

    # ========== 文件内容 ==========
    total_lines = 0

    for i, (relpath, filepath) in enumerate(files, 1):
        print(f"  [{i}/{len(files)}] {relpath.replace(os.sep, '/')}")
        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()

            _, ext = os.path.splitext(filepath)
            content = strip_comments(content, ext.lower())

            lines = content.count('\n') + 1
            total_lines += lines
            normalized = relpath.replace('\\', '/')
            file_size = os.path.getsize(filepath)
            lang = get_lang_from_ext(ext)

            heading = doc.add_heading(level=2)
            run_num = heading.add_run(f"[{i}/{len(files)}] ")
            run_num.font.color.rgb = RGBColor(100, 100, 100)
            run_num.font.size = Pt(11)
            run_path = heading.add_run(normalized)
            run_path.font.size = Pt(11)

            meta = doc.add_paragraph()
            run_meta = meta.add_run(f"{lang}  |  {lines} lines  |  {file_size:,} bytes")
            run_meta.font.size = Pt(8)
            run_meta.font.color.rgb = RGBColor(130, 130, 130)
            run_meta.font.italic = True

            code_lines = content.split('\n')
            max_lines = 2000
            truncated = False
            if len(code_lines) > max_lines:
                code_lines = code_lines[:max_lines]
                truncated = True

            chunk_size = 30
            for chunk_start in range(0, len(code_lines), chunk_size):
                chunk = code_lines[chunk_start:chunk_start + chunk_size]
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run('\n'.join(chunk))
                run.font.name = 'Consolas'
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(40, 40, 40)

            if truncated:
                p = doc.add_paragraph()
                run = p.add_run(f'\n... (truncated, showing first {max_lines} of {lines} lines) ...')
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(200, 50, 50)
                run.font.italic = True

            if i < len(files):
                sep = doc.add_paragraph()
                sep_run = sep.add_run('\u2500' * 60)
                sep_run.font.color.rgb = RGBColor(200, 200, 200)
                sep_run.font.size = Pt(8)

        except Exception as e:
            p = doc.add_paragraph()
            run = p.add_run(f"[Read failed: {e}]")
            run.font.color.rgb = RGBColor(200, 50, 50)
            run.font.bold = True

    # ========== 尾部统计 ==========
    doc.add_page_break()
    doc.add_heading('Summary', level=1)

    summary_lines = [
        f"Total files: {len(files)}",
        f"Total lines: {total_lines:,}",
        f"Export time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    ]
    for line in summary_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.save(output_file)
    output_size = os.path.getsize(output_file)

    print(f"\nDone!")
    print(f"  Files: {len(files)}")
    print(f"  Lines: {total_lines:,}")
    print(f"  Output: {output_file}")
    print(f"  Size: {output_size:,} bytes ({output_size/1024:.1f} KB)")


def main():
    parser = argparse.ArgumentParser(
        description='Export project source code to a Word document',
        epilog="""Exclude examples:
  python export_code.py --exclude output assets src/_pycache src/sfx
  python export_code.py --exclude README.md CHANGELOG.md
  python export_code.py --exclude "*.md" "*.txt"
  python export_code.py --exclude edge-extension/
  python export_code.py --exclude edge-extension/popup.js export_code.py
  python export_code.py --exclude "test_*" "*.test.js"
""",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument('directory', nargs='?', default='.')
    parser.add_argument('-o', '--output', default=None,
                        help='Output .docx file path')
    parser.add_argument('--max-size', type=int, default=200,
                        help='Max file size in KB (default: 200)')
    parser.add_argument('-e', '--exclude', nargs='+', default=[],
                        help='Files/dirs/patterns to exclude')
    parser.add_argument('--exclude-from', default=None,
                        help='Read exclude patterns from a file (one per line)')
    args = parser.parse_args()

    global MAX_FILE_SIZE, USER_EXCLUDES
    MAX_FILE_SIZE = args.max_size * 1024

    # 收集排除列表
    USER_EXCLUDES = list(args.exclude)

    # 从文件读取排除列表
    if args.exclude_from:
        try:
            with open(args.exclude_from, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        USER_EXCLUDES.append(line)
        except FileNotFoundError:
            print(f"Warning: exclude file not found: {args.exclude_from}")

    root_dir = os.path.abspath(args.directory)

    if args.output:
        output_file = args.output
    else:
        output_file = f"{os.path.basename(root_dir)}_code_export.docx"

    if not output_file.endswith('.docx'):
        output_file = output_file.rsplit('.', 1)[0] + '.docx'

    export_project(root_dir, output_file)


if __name__ == '__main__':
    main()
